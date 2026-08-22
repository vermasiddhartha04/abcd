from pathlib import Path
from typing import Any, Dict, Optional
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches


BASE_DIR = Path(__file__).resolve().parents[2]

GENERATED_DIR = (
    BASE_DIR
    / "app"
    / "services"
    / "generated_replies"
)

GENERATED_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


def _safe_filename(value: str) -> str:

    value = str(value or "reply")

    value = re.sub(
        r"[^A-Za-z0-9._-]+",
        "_",
        value,
    )

    value = value.strip(
        "._-"
    )

    return value or "reply"


def _add_paragraph(
    document: Document,
    text: str,
    bold: bool = False,
    align=None,
    size: int = 11,
):

    paragraph = document.add_paragraph()

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(
        str(text)
    )

    run.bold = bold
    run.font.size = Pt(size)

    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15

    return paragraph


def _add_heading(
    document: Document,
    text: str,
    level: int = 1,
):

    heading = document.add_heading(
        text,
        level=level,
    )

    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)

    return heading


def _format_reply_text(
    document: Document,
    reply: str,
):

    """
    Convert plain-text generated reply into
    readable DOCX paragraphs.

    Existing reply wording is preserved.
    """

    reply = reply or ""

    blocks = re.split(
        r"\n\s*\n",
        reply,
    )

    for block in blocks:

        block = block.strip()

        if not block:
            continue

        lines = block.splitlines()

        first_line = lines[0].strip()

        # --------------------------------------------------
        # SUBJECT
        # --------------------------------------------------

        if first_line.upper().startswith(
            "SUBJECT:"
        ):

            _add_paragraph(
                document,
                first_line,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=12,
            )

            if len(lines) > 1:

                remaining = "\n".join(
                    lines[1:]
                ).strip()

                if remaining:
                    _add_paragraph(
                        document,
                        remaining,
                    )

            continue

        # --------------------------------------------------
        # NUMBERED HEADINGS
        # --------------------------------------------------

        if re.match(
            r"^\d+\.\s+[A-Z]",
            first_line,
        ):

            _add_heading(
                document,
                first_line,
                level=1,
            )

            if len(lines) > 1:

                remaining = "\n".join(
                    lines[1:]
                ).strip()

                if remaining:
                    _add_paragraph(
                        document,
                        remaining,
                    )

            continue

        # --------------------------------------------------
        # ISSUE NUMBER
        # --------------------------------------------------

        if re.match(
            r"^\d+\.\s+",
            first_line,
        ):

            _add_heading(
                document,
                first_line,
                level=2,
            )

            if len(lines) > 1:

                remaining = "\n".join(
                    lines[1:]
                ).strip()

                if remaining:
                    _add_paragraph(
                        document,
                        remaining,
                    )

            continue

        # --------------------------------------------------
        # SUBHEADINGS
        # --------------------------------------------------

        if (
            first_line.endswith(":")
            and len(first_line) < 120
        ):

            _add_paragraph(
                document,
                first_line,
                bold=True,
            )

            if len(lines) > 1:

                remaining = "\n".join(
                    lines[1:]
                ).strip()

                if remaining:
                    _add_paragraph(
                        document,
                        remaining,
                    )

            continue

        # --------------------------------------------------
        # NORMAL PARAGRAPH
        # --------------------------------------------------

        _add_paragraph(
            document,
            block,
        )


def create_reply_docx(
    reply: str,
    metadata: Optional[Dict[str, Any]] = None,
    upload_id: Optional[int] = None,
    reply_id: Optional[int] = None,
) -> str:

    """
    Create a professional DOCX reply.

    Returns:
        Absolute file path.
    """

    metadata = metadata or {}

    taxpayer = (
        metadata.get(
            "taxpayer_name"
        )
        or "Taxpayer"
    )

    notice_number = (
        metadata.get(
            "notice_number"
        )
        or "SCN"
    )

    taxpayer_safe = _safe_filename(
        taxpayer
    )

    notice_safe = _safe_filename(
        notice_number
    )

    filename = (
        f"{taxpayer_safe}_"
        f"{notice_safe}_"
        f"SCN_REPLY"
    )

    if upload_id is not None:
        filename += (
            f"_upload_{upload_id}"
        )

    if reply_id is not None:
        filename += (
            f"_reply_{reply_id}"
        )

    filename += ".docx"

    output_path = (
        GENERATED_DIR
        / filename
    )

    document = Document()

    # ------------------------------------------------------
    # PAGE SETUP
    # ------------------------------------------------------

    section = document.sections[0]

    section.top_margin = Inches(
        0.8
    )

    section.bottom_margin = Inches(
        0.8
    )

    section.left_margin = Inches(
        0.9
    )

    section.right_margin = Inches(
        0.9
    )

    # ------------------------------------------------------
    # DEFAULT FONT
    # ------------------------------------------------------

    styles = document.styles

    normal_style = styles[
        "Normal"
    ]

    normal_style.font.name = (
        "Times New Roman"
    )

    normal_style.font.size = Pt(
        11
    )

    # ------------------------------------------------------
    # COVER HEADER
    # ------------------------------------------------------

    title = document.add_paragraph()

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = title.add_run(
        "REPLY TO SHOW CAUSE NOTICE"
    )

    run.bold = True
    run.font.name = (
        "Times New Roman"
    )
    run.font.size = Pt(
        16
    )

    _add_paragraph(
        document,
        f"Taxpayer: {taxpayer}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_paragraph(
        document,
        f"Notice Number: {notice_number}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    gstin = metadata.get(
        "gstin"
    )

    if gstin:
        _add_paragraph(
            document,
            f"GSTIN: {gstin}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    financial_year = metadata.get(
        "financial_year"
    )

    if financial_year:
        _add_paragraph(
            document,
            f"Financial Year: {financial_year}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    document.add_page_break()

    # ------------------------------------------------------
    # REPLY CONTENT
    # ------------------------------------------------------

    _format_reply_text(
        document,
        reply,
    )

    # ------------------------------------------------------
    # FOOTER
    # ------------------------------------------------------

    for section in document.sections:

        footer = section.footer

        paragraph = footer.paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            "Generated GST Litigation AI Reply"
        )

        run.font.size = Pt(
            8
        )

    # ------------------------------------------------------
    # SAVE
    # ------------------------------------------------------

    document.save(
        str(output_path)
    )

    return str(
        output_path
    )


def create_reply_document(
    reply: str,
    metadata: Optional[Dict[str, Any]] = None,
    upload_id: Optional[int] = None,
    reply_id: Optional[int] = None,
) -> str:

    """
    Backward-compatible alias.
    """

    return create_reply_docx(
        reply=reply,
        metadata=metadata,
        upload_id=upload_id,
        reply_id=reply_id,
    )
